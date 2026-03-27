"""
DOCX document processor using python-docx.
"""
import io

from processors.base import BaseProcessor


class DocxProcessor(BaseProcessor):
    """Extracts plain text from DOCX files using python-docx."""

    def process(self, file_bytes: bytes) -> str:
        """Extract text from a DOCX file.

        Extracts text from paragraphs and table cells, preserving
        a logical reading order.

        Args:
            file_bytes: Raw bytes of the DOCX document.

        Returns:
            Extracted plain text from paragraphs and tables.

        Raises:
            ImportError: If python-docx is not installed.
            Exception: If the DOCX cannot be parsed.
        """
        try:
            import docx
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install it with: pip install python-docx"
            ) from exc

        document = docx.Document(io.BytesIO(file_bytes))
        text_parts: list[str] = []

        # Extract paragraph text
        for paragraph in document.paragraphs:
            stripped = paragraph.text.strip()
            if stripped:
                text_parts.append(stripped)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_texts:
                    text_parts.append(" | ".join(row_texts))

        return "\n\n".join(text_parts)
